from datasets import load_dataset
import csv
import base64
import os
import json
import time
from huggingface_hub import HfApi
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pdf2image import convert_from_path
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.oxml.ns import qn
from PIL import Image
import requests
from itertools import islice
import subprocess

leftover_words = []
# Initialize Huggingface API
api = HfApi(token="YOUR_API_TOKEN")
repo_id = "YOUR_REPO" 

DATASET_NAME = "DATASET_NAME"
OUTPUT_FOLDER = "./directory"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Path to the JSON state file
STATE_FILE = "processing_state.json"

# Function to load the processing state from the JSON file
def load_state():
    if os.path.exists(STATE_FILE):
        with open(STATE_FILE, "r") as f:
            state = json.load(f)
            print(f"Loaded state: {state}")  # Debugging statement
            return state
    print("State file not found, initializing with default values.")  # Debugging
    return {"article_index": 0, "batch_index": 0, "font_name_index": 0}

# Function to save the processing state to the JSON file
def save_state(index, batch_index, font_name_index):
    state = {"article_index": index, "batch_index": batch_index, "font_name_index": font_name_index}
    try:
        with open(STATE_FILE, "w") as f:
            json.dump(state, f)
            print(f"Saved state: {state}")  # Debugging
    except Exception as e:
        print(f"Error saving state: {e}")

def split_text_into_chunks(text, max_words=405, min_words=50):
    words = text.split()
    chunks = []
    global leftover_words
    # Merge with leftover words from the previous call
    if leftover_words:
        words = leftover_words + words
        leftover_words = []  # Clear the buffer
    
    temp_chunk = []
    
    for word in words:
        temp_chunk.append(word)

        if len(temp_chunk) >= max_words:  # When max_words is reached
            chunks.append(" ".join(temp_chunk))
            temp_chunk = []

    # Handling remaining words
    if temp_chunk:
        if len(temp_chunk) < min_words:
            # Save as leftover to merge with the next incoming text
            leftover_words = temp_chunk
        else:
            chunks.append(" ".join(temp_chunk))

    return chunks
def create_docx(text, file_name, font_name):
    try:
        doc = Document()
        section = doc.sections[0]
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.gutter = Inches(0.2)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

        for line in text.split('\n'):
            if not line.strip():
                doc.add_paragraph("")
                continue
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.line_spacing = 1.5

            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            paragraph._element.get_or_add_pPr().append(bidi)

            run = paragraph.add_run(line)
            rPr = run._element.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                run._element.insert(0, rPr)

            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), font_name)
            rFonts.set(qn("w:hAnsi"), font_name)
            rFonts.set(qn("w:cs"), font_name)
            rPr.append(rFonts)
            
            if font_name == "Sakkal Majalla":
                font_size = 14
            elif font_name == "Amiri":
                font_size = 12
            elif font_name == "Arial":
                font_size = 12
            elif font_name == "Calibri":
                font_size = 12
            elif font_name == "Scheherazade New":
                font_size = 12

            size_sz = OxmlElement("w:sz")
            size_sz.set(qn("w:val"), str(font_size * 2))
            rPr.append(size_sz)
            size_szCs = OxmlElement("w:szCs")
            size_szCs.set(qn("w:val"), str(font_size * 2))
            rPr.append(size_szCs)

        doc.save(file_name)
    except Exception as e:
        print(f"Error in create_docx: {e}")

def convert_docx_to_pdf(docx_file, pdf_file):
    # os.system(f"libreoffice --headless --convert-to pdf {docx_file} --outdir {os.path.dirname(pdf_file)}")
    # if not os.path.exists(pdf_file):
    #     raise FileNotFoundError(f"Failed to convert {docx_file} to PDF.")
    with open(os.devnull, 'w') as fnull:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", docx_file, "--outdir", os.path.dirname(pdf_file)],
            stdout=fnull, stderr=subprocess.PIPE  # Capture errors only
        )
    if result.returncode != 0:
        print(f"Failed to convert {docx_file} to PDF. Error: {result.stderr.decode()}")

    

def convert_pdf_to_image(pdf_file):
    pages = convert_from_path(pdf_file, dpi=300)
    widths, heights = zip(*(page.size for page in pages))
    total_height = sum(heights)
    max_width = max(widths)

    combined_image = Image.new('RGB', (max_width, total_height), (255, 255, 255))
    y_offset = 0
    for page in pages:
        combined_image.paste(page, (0, y_offset))
        y_offset += page.size[1]

    output_image_path = os.path.join(OUTPUT_FOLDER, f"{os.path.splitext(os.path.basename(pdf_file))[0]}.png")
    combined_image.save(output_image_path, "PNG")

    return output_image_path

def convert_image_to_base64(image_path):
    if not os.path.exists(image_path):
        raise FileNotFoundError(f"Image file not found: {image_path}")
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode('utf-8')

# Function to handle API uploads with retries
def upload_file_with_retries(api, path_or_fileobj,subset_name, batch_csv_file, repo_id, retries=5, delay=60):
    for attempt in range(retries):
        try:
            api.upload_file(
                path_or_fileobj=path_or_fileobj,
                path_in_repo=f"{subset_name.replace(' ', '_')}/{batch_csv_file}", #path_in_repo,
                repo_id=repo_id,
                repo_type="dataset",
                commit_message=f"Add file {batch_csv_file} to {subset_name.replace(' ', '_')}",
            )
            print(f"Successfully uploaded: {batch_csv_file}")
            return
        except requests.exceptions.ConnectionError as e:
            print(f"ConnectionError on attempt {attempt + 1}/{retries}: {e}")
            if attempt < retries - 1:
                time.sleep(delay)
            else:
                print(f"Failed to upload {batch_csv_file} after {retries} attempts.")
                raise

# Load dataset and state
ds = load_dataset("USE YOUR HUGGINGFACE DATASET PATH", streaming=True)
state = load_state()
article_index = state["article_index"]
batch_index = state["batch_index"]
font_name_index = state["font_name_index"]

batch_size = 1000
current_batch = []
font_name_list = ["Sakkal Majalla", "Amiri", "Arial", "Calibri", "Scheherazade New"]
# Process dataset

for font_name in font_name_list[font_name_index:]:
    print("font_name:", font_name)
    for index, row in enumerate(islice(ds["train"], article_index, None), start=article_index):

        try:
            max_words =0
            if font_name == "Sakkal Majalla":
                max_words = 300
            elif font_name == "Arial":
                max_words = 500
            elif font_name == "Calibri":
                max_words = 500
            elif font_name == "Amiri":
                max_words = 300
            elif font_name == "Scheherazade New":
                max_words = 250

            chunks = split_text_into_chunks(row['text'], max_words)

            for chunk_idx, chunk in enumerate(chunks):
                base_name = f"dataset_{DATASET_NAME}_font_{font_name.replace(' ', '_')}_article_{index+1}_part_{chunk_idx+1}"
                docx_file = os.path.join(OUTPUT_FOLDER, f"{base_name}.docx")
                pdf_file = os.path.join(OUTPUT_FOLDER, f"{base_name}.pdf")

                create_docx(chunk, docx_file, font_name)
                convert_docx_to_pdf(docx_file, pdf_file)
                image_path = convert_pdf_to_image(pdf_file)
                img_base64 = convert_image_to_base64(image_path)
                current_batch.append([os.path.basename(image_path), chunk, font_name, img_base64])

                os.remove(docx_file)
                os.remove(pdf_file)
                os.remove(image_path)

            if (index + 1) % batch_size == 0:
                batch_csv_file = f"dataset_{DATASET_NAME}_font_{font_name.replace(' ', '_')}_batch_{batch_index}.csv"
                batch_csv_path = os.path.join(OUTPUT_FOLDER, batch_csv_file)

                with open(batch_csv_path, mode='w', newline='', encoding='utf-8') as csvfile:
                    csvwriter = csv.writer(csvfile)
                    csvwriter.writerow(["image_name", "chunk", "font_name", "image_base64"])
                    csvwriter.writerows(current_batch)

                upload_file_with_retries(api, batch_csv_path, font_name, batch_csv_file, repo_id)
                os.remove(batch_csv_path)
                current_batch = []
                batch_index += 1
                save_state(index + 1, batch_index, font_name_list.index(font_name))

        except Exception as e:
            print(f"Error processing index {index}: {e}")
         

    if current_batch:
        batch_csv_file = f"dataset_{DATASET_NAME}_font_{font_name.replace(' ', '_')}_batch_{batch_index}.csv"
        batch_csv_path = os.path.join(OUTPUT_FOLDER, batch_csv_file)

        with open(batch_csv_path, mode='w', newline='', encoding='utf-8') as csvfile:
            csvwriter = csv.writer(csvfile)
            csvwriter.writerow(["image_name", "chunk", "font_name", "image_base64"])
            csvwriter.writerows(current_batch)

        upload_file_with_retries(api, batch_csv_path,font_name, batch_csv_file, repo_id)
        os.remove(batch_csv_path)

        save_state(0, 0, font_name_list.index(font_name)+1)
    
    current_batch = []
    batch_index = 0
    article_index = 0

print("Processing completed.")
